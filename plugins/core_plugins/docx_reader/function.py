from typing import Any, Dict, List

from services.path_utils import resolve_user_file


def docx_reader(file_path: str, include_tables: bool = False) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ImportError("docx_reader requires the 'python-docx' package.") from exc

    doc_path = resolve_user_file(file_path, expected_extensions=(".docx",))
    document = Document(doc_path)

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]

    tables: List[str] = []
    if include_tables:
        for table in document.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables.append(" | ".join(row_text))

    payload: Dict[str, Any] = {"paragraphs": paragraphs}
    if include_tables:
        payload["tables"] = tables
    payload["meta"] = {"path": str(doc_path), "paragraph_count": len(paragraphs), "tables_included": include_tables}
    return payload
